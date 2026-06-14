"""Testes de extração docx (tabelas de insumos)."""

from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from rag.extract import _extract_doc_header, _extract_docx  # noqa: E402


class ExtractDocxTests(unittest.TestCase):
    def test_tables_included_in_extraction(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "planejamento.docx"
            doc = Document()
            doc.add_paragraph("Planejamento do ensaio ELISA")
            table = doc.add_table(rows=2, cols=4)
            hdr = table.rows[0].cells
            hdr[0].text = "Nome"
            hdr[1].text = "Fabricante ou código"
            hdr[2].text = "Lote"
            hdr[3].text = "Validade"
            row = table.rows[1].cells
            row[0].text = "Tampão de amostra"
            row[1].text = "FAB-001"
            row[2].text = "L2024-09"
            row[3].text = "31/12/2026"
            doc.save(str(path))

            outcome = _extract_docx(path, max_chars_total=500_000)
            self.assertTrue(outcome.ok, outcome.detail)
            self.assertIn("Planejamento", outcome.text)
            self.assertIn("Tampão de amostra", outcome.text)
            self.assertIn("31/12/2026", outcome.text)
            self.assertIn("tabela", outcome.detail.lower())

    def test_doc_header_extracted_from_planning_dates(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "ensaio.docx"
            doc = Document()
            doc.add_paragraph("ENSAIO OTIMIZACAO TEMPO DE SENSIBILIZACAO")
            doc.add_paragraph("Planejamento: 09/02/2026")
            doc.add_paragraph("Execução: 09/02/2026")
            doc.add_paragraph("Objetivo")
            doc.save(str(path))

            outcome = _extract_docx(path, max_chars_total=500_000)
            self.assertIsNotNone(outcome.doc_header)
            assert outcome.doc_header is not None
            self.assertIn("ENSAIO OTIMIZACAO TEMPO DE SENSIBILIZACAO", outcome.doc_header)
            self.assertIn("Planejamento: 09/02/2026", outcome.doc_header)
            self.assertIn("Execução: 09/02/2026", outcome.doc_header)

    def test_doc_header_none_without_planning_dates(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sem_data.docx"
            doc = Document()
            doc.add_paragraph("Documento qualquer sem cabeçalho de ensaio")
            doc.add_paragraph("Texto normal de conteúdo.")
            doc.save(str(path))

            outcome = _extract_docx(path, max_chars_total=500_000)
            self.assertIsNone(outcome.doc_header)

    def test_extract_doc_header_helper_direct(self) -> None:
        text = (
            "ENSAIO X\n"
            "Planejamento: 10/02/2026\n"
            "Execução: 12/02/2026\n"
            "Objetivo\n"
        )
        header = _extract_doc_header(text)
        self.assertEqual(header, "[Ensaio: ENSAIO X | Planejamento: 10/02/2026 | Execução: 12/02/2026]")

    def test_nested_table_in_cell(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "nested.docx"
            doc = Document()
            outer = doc.add_table(rows=1, cols=1)
            cell = outer.rows[0].cells[0]
            cell.text = "Reagente principal"
            inner = cell.add_table(rows=2, cols=2)
            inner.rows[0].cells[0].text = "Lote"
            inner.rows[0].cells[1].text = "Validade"
            inner.rows[1].cells[0].text = "L-99"
            inner.rows[1].cells[1].text = "01/06/2027"
            doc.save(str(path))

            outcome = _extract_docx(path, max_chars_total=500_000)
            self.assertTrue(outcome.ok, outcome.detail)
            self.assertIn("Reagente principal", outcome.text)
            self.assertIn("L-99", outcome.text)
            self.assertIn("01/06/2027", outcome.text)


if __name__ == "__main__":
    unittest.main()
